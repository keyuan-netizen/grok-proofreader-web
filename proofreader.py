import json, csv
from pathlib import Path
from docx import Document
import requests
from typing import List, Dict
import logging

DEFAULT_API_URL = "https://api.x.ai/v1/chat/completions"
DEFAULT_MODEL = "grok-3"

PROMPT = (
    "Proofread and return ONLY valid JSON: "
    '{{"corrections": [{{"original": "", "suggested": "", "reason": ""}}], '
    '"summary": "..."}} Text: """{}"""'
)

logger = logging.getLogger("proofreader")
if not logger.handlers:
    logging.basicConfig(level=logging.INFO)

def extract_text(path: Path) -> str:
    return "\n".join(p.text for p in Document(path).paragraphs if p.text.strip())

def call_grok(text: str, api_key: str, system_prompt: str, api_url=DEFAULT_API_URL, model=DEFAULT_MODEL):
    user_prompt = PROMPT.format(text)
    logger.info(
        "Dispatching Grok request with system prompt:\n%s\nUser prompt:\n%s",
        system_prompt,
        user_prompt,
    )
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
    }
    resp = requests.post(
        api_url,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json=payload,
        timeout=60
    )
    resp.raise_for_status()
    data = resp.json()
    content = data["choices"][0]["message"]["content"]
    return json.loads(content)

def build_table(doc, corrections):
    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    for i, label in enumerate(["Original", "Suggested", "Reason"]):
        table.rows[0].cells[i].text = label
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for c in corrections:
        row = table.add_row().cells
        row[0].text = c.get("original", "")
        row[1].text = c.get("suggested", "")
        row[2].text = c.get("reason", "")

def save_reports(results: List[Dict], out_dir: Path):
    out_dir.mkdir(parents=True, exist_ok=True)
    
    # JSON + CSV
    (out_dir / "results.json").write_text(json.dumps(results, indent=2), encoding="utf-8")
    with (out_dir / "summary.csv").open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["filename", "status", "correction_count", "summary"])
        for r in results:
            d = r["api_result"]["data"]
            w.writerow([r["filename"], "proofread", len(d.get("corrections", [])), d.get("summary", "")])
    
    # Master DOCX
    master = Document()
    master.add_heading("Grok Proofreading Report", 0)
    for r in results:
        d = r["api_result"]["data"]
        master.add_heading(r["filename"], 1)
        master.add_paragraph(f"Characters: {r['char_count']}")
        master.add_paragraph(f"Summary: {d.get('summary', '')}")
        build_table(master, d.get("corrections", []))
        master.add_page_break()
    master.save(out_dir / "PROOFREADING_REPORT.docx")
    
    # Per-file DOCX
    for r in results:
        d = r["api_result"]["data"]
        doc = Document()
        doc.add_heading(f"Proofreading: {r['filename']}", 0)
        doc.add_paragraph(f"Summary: {d.get('summary', '')}")
        build_table(doc, d.get("corrections", []))
        doc.save(out_dir / f"{Path(r['filename']).stem}_PROOFREAD.docx")
